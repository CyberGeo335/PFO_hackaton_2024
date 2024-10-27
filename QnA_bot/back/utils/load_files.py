# app/utils/load_files.py

import os
from docx import Document as DocxDocument
from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Функция для разбиения файлов на чанки
def file_to_chunks_with_splitter(file_name,
                                 sep,
                                 chunk_size,
                                 chunk_overlap):
    """
        Разбивает файл на чанки текста заданного размера и перекрытия.

        Параметры:
            file_name (str): Полный путь к файлу.
            sep (str): Разделитель для объединения текста в случае таблиц.
            chunk_size (int): Размер каждого чанка текста.
            chunk_overlap (int): Размер перекрытия между чанками.

        Возвращает:
            list: Список текстовых чанков, полученных из файла.

        Поддерживаемые форматы:
            - txt
            - docx
            - pdf

        Исключения:
            ValueError: Выдается, если формат файла не поддерживается.
        """
    file_ext = file_name.split('.')[-1]
    file_path = file_name
    overall_chunks = []

    if file_ext == 'txt':
        loader = TextLoader(file_path, encoding='utf-8')
        text_content = loader.load()
    elif file_ext == 'docx':
        loader = DocxDocument(file_path)
        text_content = []
        for para in loader.paragraphs:
            if para.text.strip():
                text_content.append(para.text.strip())
        for table in loader.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                text_content.append(sep.join(row_data))
    elif file_ext == 'pdf':
        loader = PyPDFLoader(file_path)
        text_content = [page.extract_text() for page in loader.load() if page.extract_text() is not None]
    else:
        raise ValueError("Unsupported file format")

    text_splitter = RecursiveCharacterTextSplitter(
        separators=sep,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        is_separator_regex=False,
        add_start_index=False
    )

    for content in text_content:
        if content.strip():
            chunks = text_splitter.split_text(content)
            overall_chunks.extend([chunk.strip() for chunk in chunks if chunk.strip()])

    return overall_chunks

# Функция для загрузки и обработки проекта с вложенными файлами
def load_project_files(project_dir,
                       sep='\n',
                       chunk_size=2048,
                       chunk_overlap=128):
    """
        Загружает и разбивает все файлы в директории проекта на чанки,
        преобразуя их в объекты Document для обработки.

        Параметры:
            project_dir (str): Путь к директории проекта.
            sep (str): Разделитель для текста в случае таблиц.
            chunk_size (int): Размер каждого чанка.
            chunk_overlap (int): Размер перекрытия между чанками.

        Возвращает:
            list: Список объектов Document, содержащих текстовые чанки и метаданные.
    """
    documents = []
    for root, _, files in os.walk(project_dir):
        for file in files:
            if file.startswith('.'):
                continue
            file_path = os.path.join(root, file)
            try:
                chunks = file_to_chunks_with_splitter(file_path, sep, chunk_size, chunk_overlap)
                for chunk in chunks:
                    metadata = {"название документа": file_path}
                    documents.append(Document(
                        text=chunk,
                        metadata=metadata,
                        excluded_embed_metadata_keys=["название документа"]
                    ))
            except Exception as e:
                print(f"Ошибка при обработке файла {file_path}: {e}")
    return documents
